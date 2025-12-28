"""
Resume Parser - Extract text from PDF, Word, and text files
"""
import logging
from io import BytesIO
from typing import Optional
import PyPDF2
from docx import Document


class ResumeParser:
    """Parser for different resume file formats"""
    
    def __init__(self):
        """Initialize resume parser"""
        logging.info("Resume Parser initialized")
    
    def parse(self, file_obj, file_type: str) -> Optional[str]:
        """
        Parse resume file and extract text
        
        Args:
            file_obj: File object or bytes
            file_type: File extension (pdf, docx, txt)
            
        Returns:
            Extracted text or None if error
        """
        file_type = file_type.lower().strip('.')
        
        try:
            if file_type == 'pdf':
                return self._parse_pdf(file_obj)
            elif file_type in ['docx', 'doc']:
                return self._parse_docx(file_obj)
            elif file_type == 'txt':
                return self._parse_txt(file_obj)
            else:
                logging.error(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logging.error(f"Error parsing {file_type} file: {str(e)}")
            return None
    
    def _parse_pdf(self, file_obj) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(file_obj)
            
            # Extract text from all pages
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
            
            if text.strip():
                logging.info(f"Successfully extracted text from PDF ({len(text)} characters)")
                return text.strip()
            else:
                logging.warning("PDF text extraction resulted in empty text")
                return None
                
        except Exception as e:
            logging.error(f"Error parsing PDF: {str(e)}")
            return None
    
    def _parse_docx(self, file_obj) -> Optional[str]:
        """Extract text from Word document"""
        try:
            # Create a Document object
            doc = Document(file_obj)
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                logging.info(f"Successfully extracted text from DOCX ({len(text)} characters)")
                return text.strip()
            else:
                logging.warning("DOCX text extraction resulted in empty text")
                return None
                
        except Exception as e:
            logging.error(f"Error parsing DOCX: {str(e)}")
            return None
    
    def _parse_txt(self, file_obj) -> Optional[str]:
        """Extract text from plain text file"""
        try:
            # Read text file
            if hasattr(file_obj, 'read'):
                text = file_obj.read()
                if isinstance(text, bytes):
                    text = text.decode('utf-8')
            else:
                text = str(file_obj)
            
            if text.strip():
                logging.info(f"Successfully extracted text from TXT ({len(text)} characters)")
                return text.strip()
            else:
                logging.warning("TXT file is empty")
                return None
                
        except Exception as e:
            logging.error(f"Error parsing TXT: {str(e)}")
            return None
    
    def validate_file_size(self, file_obj, max_size_mb: int = 10) -> bool:
        """
        Validate that file size is within limits
        
        Args:
            file_obj: File object
            max_size_mb: Maximum file size in MB
            
        Returns:
            True if valid, False otherwise
        """
        try:
            # Get file size
            if hasattr(file_obj, 'size'):
                size_bytes = file_obj.size
            elif hasattr(file_obj, 'seek') and hasattr(file_obj, 'tell'):
                current_pos = file_obj.tell()
                file_obj.seek(0, 2)  # Seek to end
                size_bytes = file_obj.tell()
                file_obj.seek(current_pos)  # Restore position
            else:
                logging.warning("Cannot determine file size")
                return True  # Allow if can't determine
            
            max_bytes = max_size_mb * 1024 * 1024
            
            if size_bytes > max_bytes:
                logging.warning(f"File size ({size_bytes} bytes) exceeds limit ({max_bytes} bytes)")
                return False
            
            return True
            
        except Exception as e:
            logging.error(f"Error validating file size: {str(e)}")
            return True  # Allow if error
    
    def extract_sections(self, text: str) -> dict:
        """
        Extract common resume sections (basic implementation)
        
        Args:
            text: Resume text
            
        Returns:
            Dict with identified sections
        """
        sections = {
            'full_text': text,
            'has_contact': False,
            'has_experience': False,
            'has_education': False,
            'has_skills': False
        }
        
        text_lower = text.lower()
        
        # Check for common section headers
        if any(keyword in text_lower for keyword in ['email', 'phone', '@', 'linkedin']):
            sections['has_contact'] = True
        
        if any(keyword in text_lower for keyword in ['experience', 'work history', 'employment']):
            sections['has_experience'] = True
        
        if any(keyword in text_lower for keyword in ['education', 'degree', 'university', 'college']):
            sections['has_education'] = True
        
        if any(keyword in text_lower for keyword in ['skills', 'technologies', 'proficient']):
            sections['has_skills'] = True
        
        logging.info(f"Resume sections identified: {sections}")
        return sections
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join with single newlines
        cleaned = '\n'.join(lines)
        
        # Remove multiple spaces
        import re
        cleaned = re.sub(r' +', ' ', cleaned)
        
        return cleaned